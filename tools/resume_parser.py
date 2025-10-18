"""Resume parsing tools for extracting information from candidate resumes."""

import io
import re
from typing import Optional, List, Dict, Any
from pathlib import Path
from pydantic import BaseModel, Field

import PyPDF2
import pdfplumber
import docx
import spacy


class ParsedResume(BaseModel):
    """Structured resume data."""
    
    raw_text: str = Field(description="Raw extracted text from resume")
    name: Optional[str] = Field(default=None, description="Candidate name")
    email: Optional[str] = Field(default=None, description="Email address")
    phone: Optional[str] = Field(default=None, description="Phone number")
    skills: List[str] = Field(default_factory=list, description="Extracted skills")
    experience: List[Dict[str, Any]] = Field(
        default_factory=list,
        description="Work experience entries"
    )
    education: List[Dict[str, Any]] = Field(
        default_factory=list,
        description="Education entries"
    )
    summary: Optional[str] = Field(default=None, description="Professional summary")


class ResumeParser:
    """Parser for extracting structured information from resumes."""
    
    def __init__(self):
        """Initialize the resume parser."""
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            # Model not downloaded yet
            self.nlp = None
    
    def parse_file(self, file_path: str) -> ParsedResume:
        """
        Parse a resume file and extract structured information.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            ParsedResume object with extracted information
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        # Extract text based on file type
        if file_path.suffix.lower() == '.pdf':
            text = self._extract_pdf_text(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            text = self._extract_docx_text(file_path)
        elif file_path.suffix.lower() == '.txt':
            text = self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        # Parse the extracted text
        return self._parse_text(text)
    
    def parse_bytes(self, file_bytes: bytes, file_extension: str) -> ParsedResume:
        """
        Parse resume from bytes.
        
        Args:
            file_bytes: Resume file as bytes
            file_extension: File extension (e.g., '.pdf', '.docx')
            
        Returns:
            ParsedResume object with extracted information
        """
        if file_extension.lower() == '.pdf':
            text = self._extract_pdf_bytes(file_bytes)
        elif file_extension.lower() in ['.docx', '.doc']:
            text = self._extract_docx_bytes(file_bytes)
        elif file_extension.lower() == '.txt':
            text = file_bytes.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        return self._parse_text(text)
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = ""
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text
        except Exception:
            pass
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise ValueError(f"Failed to extract PDF text: {str(e)}")
        
        return text
    
    def _extract_pdf_bytes(self, file_bytes: bytes) -> str:
        """Extract text from PDF bytes."""
        text = ""
        
        # Try pdfplumber first
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text
        except Exception:
            pass
        
        # Fallback to PyPDF2
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            raise ValueError(f"Failed to extract PDF text: {str(e)}")
        
        return text
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX text: {str(e)}")
    
    def _extract_docx_bytes(self, file_bytes: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX text: {str(e)}")
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise ValueError(f"Failed to read text file: {str(e)}")
    
    def _parse_text(self, text: str) -> ParsedResume:
        """
        Parse extracted text and structure the information.
        
        Args:
            text: Raw text from resume
            
        Returns:
            ParsedResume with structured data
        """
        # Extract email
        email = self._extract_email(text)
        
        # Extract phone
        phone = self._extract_phone(text)
        
        # Extract name (basic heuristic - first line or first person entity)
        name = self._extract_name(text)
        
        # Extract skills
        skills = self._extract_skills(text)
        
        return ParsedResume(
            raw_text=text,
            name=name,
            email=email,
            phone=phone,
            skills=skills,
            experience=[],  # Would need more sophisticated parsing
            education=[],   # Would need more sophisticated parsing
        )
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text."""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)
        return matches[0] if matches else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text."""
        # Various phone number patterns
        patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            if matches:
                return matches[0]
        
        return None
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name from text."""
        # Try to get first line (common resume format)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if lines:
            first_line = lines[0]
            # Check if it looks like a name (not too long, no special chars)
            if len(first_line.split()) <= 4 and not re.search(r'[^a-zA-Z\s]', first_line):
                return first_line
        
        # Try NLP entity extraction if available
        if self.nlp:
            doc = self.nlp(text[:500])  # Check first 500 chars
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    return ent.text
        
        return None
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text."""
        # Common technical skills (simplified list)
        common_skills = [
            'python', 'java', 'javascript', 'typescript', 'react', 'angular', 'vue',
            'node.js', 'django', 'flask', 'fastapi', 'sql', 'postgresql', 'mysql',
            'mongodb', 'redis', 'docker', 'kubernetes', 'aws', 'azure', 'gcp',
            'machine learning', 'deep learning', 'nlp', 'computer vision',
            'git', 'agile', 'scrum', 'leadership', 'communication', 'teamwork'
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in common_skills:
            if skill in text_lower:
                found_skills.append(skill)
        
        return found_skills
